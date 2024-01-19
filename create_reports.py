import os
from copy import deepcopy
from datetime import datetime as dt
from dotenv import load_dotenv

import numpy as np
import pandas as pd
from sqlalchemy.engine import create_engine

import docx

load_dotenv()


dest_dir = r"Z:\EP&RS\Customer Satisfaction\Complaint Management\MPSC Case Information\MPSC Templates\MPSC Templates 2023"

# --- Engines

print("Creating Databricks engine...")
adb_engine = create_engine(
    f"databricks://token:{os.environ['DATABRICKS_TOKEN']}@{os.environ['DATABRICKS_HOST']}:443/edw?http_path={os.environ['DATABRICKS_HTTP_PATH']}")

print("Creating Oracle engine...")
ora_engine = create_engine(
    f'oracle+cx_oracle://{os.environ["ORACLE_USER"]}:{os.environ["ORACLE_PASS"]}@{os.environ["ORACLE_DBDSN"]}')

# --- Reading Data from Sources

print("Reading all sheets from Excel file...")
xl = pd.read_excel("Example for Safwaan.xlsx", sheet_name=None)
surge_2023 = pd.read_excel("2023 Surge Circuits for Dashboard.xlsx", sheet_name="Final")

tt_db = xl["TT Circuits DB"]
tt_2023 = xl["TT 2023 Plan Work"]
tt_2022 = xl["TT 2022 Carry Over"]
tt_2022.columns = tt_2022.iloc[0]
tt_2022.drop(tt_2022.index[:2], inplace=True)
ptm = xl["PTM "]
ce = xl["CE Circuits"]

print("Fetching case data...")
with open("get_cases.sql", "r") as f:
    case_query = f.read()
case_df = pd.read_sql(case_query, con=ora_engine)

premises = case_df["premise"][~pd.isna(case_df["premise"])].astype(
    np.uint64).unique().tolist()

with open("special-cases.txt", "r") as sc:
    special_cases = dict([[int(i) for i in line.split()] for line in sc])

premises.extend(set(special_cases.values()))

print("Fetching ADB AMI data...")
with open("get_adb_ami_data.sql", "r") as f:
    adb_ami_query = f.read().format(
        ",".join("%s" for _ in premises))
adb_ami_df = pd.read_sql(adb_ami_query, con=adb_engine, params=premises)
adb_ami_df["aoe_premise"] = adb_ami_df["aoe_premise"].astype(np.uint64)

print("Fetching on-prem AMI data...")
with open("get_ora_ami_data.sql", "r") as f:
    ora_ami_query = f.read().format(
        ",".join(f":{i}" for i, _ in enumerate(premises, 1)))
ora_ami_df = pd.read_sql(ora_ami_query, con=ora_engine, params=premises)

print("Fetching ADB customer data...")
with open("get_customer_data.sql", "r") as f:
    adb_cust_query = f.read().format(",".join(f"%s" for _ in premises))
cust_df = pd.read_sql(adb_cust_query, con=adb_engine, params=premises)
cust_df["premise_number"] = cust_df["premise_number"].astype(np.uint64)

adb_ami_df["aoe_down_date"] = adb_ami_df["aoe_down_date"].dt.tz_localize(None)
adb_ami_df["aoe_up_date"] = adb_ami_df["aoe_up_date"].dt.tz_localize(None)
adb_ami_df["Sustained/Momentary"] = "Sustained"
ora_ami_df["Sustained/Momentary"] = "Momentary"

out_hist_df = pd.concat([adb_ami_df, ora_ami_df]).reset_index(
    drop=True).replace(np.nan, None).fillna("null").sort_values(["Sustained/Momentary", "aoe_down_date"], ascending=[False, False])
out_hist_df["aoe_duration_minutes"] = out_hist_df["aoe_duration_minutes"].apply(
    lambda v: "{:02}:{:02}".format(*divmod(int(v), 60)))
# out_hist_df.drop(columns="aoe_creation_date", inplace=True)

# --- Creating Reports

template_doc = docx.Document("Template.docx")
template_fo_doc = docx.Document("Template-FO-noemail.docx")

cur_date = dt.now().date()
cur_dest = f"{dest_dir}{os.sep}Template Automation{os.sep}{cur_date}"
os.makedirs(cur_dest, exist_ok=True)

print(len(case_df))
for i, (case, create_date, premise, frequent) in case_df.iterrows():
    premise = special_cases.get(case, premise)
    print("new", case, create_date, premise)
    if pd.isna(premise):
        print(f"No premise for case {case}")
        continue

    doc = deepcopy(template_fo_doc if frequent else template_doc)
    doc.paragraphs[0].text = f"Case {case} {dt.strptime(create_date, '%m/%d/%Y').strftime('%b %d, %Y')}"

    cust = cust_df[["customer_name", "service_address"]
                   ][cust_df["premise_number"] == premise]
    circuit = adb_ami_df["latest_circuit"][adb_ami_df["aoe_premise"]
                                           == premise].iloc[:1]
    if circuit.size:
        [circuit] = circuit
        circuit_nospace = circuit[:-4].strip() + circuit[-4:]
    else:
        circuit = circuit_nospace = "No Data"

    longest_sustained = adb_ami_df["aoe_duration_minutes"].astype(
        float).where(adb_ami_df["aoe_premise"] == premise).max()

    num_sustained = (adb_ami_df["aoe_premise"] == premise).sum()
    num_momentary = (ora_ami_df["aoe_premise"] == premise).sum()

    doc.tables[0].cell(0, 1).paragraphs[0].add_run(str(case))
    doc.tables[0].cell(1, 1).paragraphs[0].add_run(
        *cust["customer_name"].values)
    doc.tables[0].cell(2, 1).paragraphs[0].add_run(
        *cust["service_address"].values)

    doc.tables[1].cell(0, 2).paragraphs[0].add_run(create_date)
    doc.tables[1].cell(1, 2).paragraphs[0].add_run(circuit)
    # doc.tables[1].cell(2, 2).paragraphs[0].add_run(
    #     "Yes" if longest_sustained > 5760 else "No")

    for p in doc.tables[1].cell(4, 2).paragraphs:
        for r in p.runs:
            r.text = r.text.replace("<NUM_SUSTAINED>", str(num_sustained)).replace(
                "<NUM_MOMENTARY>", str(num_momentary))

    doc.tables[1].cell(4, 2).add_paragraph()
    if circuit == "No Data":
        last_trim = "No Data"
        doc.tables[1].cell(4, 2).paragraphs[-1].add_run(f"Last Trim: No Data")
    else:
        [last_trim] = tt_db["LAST_TRIM"][tt_db["CIRCUIT"]
                                         == circuit].values
        doc.tables[1].cell(4, 2).paragraphs[-1].add_run(
            f"Last Trim: {'No Data' if pd.isna(last_trim) else int(last_trim)}")

    planned_2023 = False if circuit == "No Data" else (
        tt_2023["Circuit"] == circuit_nospace).sum()

    doc.tables[1].cell(4, 2).add_paragraph()
    doc.tables[1].cell(
        4, 2).paragraphs[-1].add_run(f"TT 2023 Planned Work: {'Yes' if planned_2023 else 'No'}")

    carryover_2022 = False if circuit == "No Data" else (
        tt_2022["Circuit"] == circuit_nospace).sum()

    doc.tables[1].cell(4, 2).add_paragraph()
    doc.tables[1].cell(
        4, 2).paragraphs[-1].add_run(f"TT 2022 Carry Over: {'Yes' if carryover_2022 else 'No'}")

    doc.tables[1].cell(4, 2).add_paragraph()
    if circuit == "No Data" or not (x := ptm["Construction Completion Date"][ptm["Circuit"] == circuit_nospace].sort_values(ascending=False).iloc[:1]).size:
        doc.tables[1].cell(4, 2).paragraphs[-1].add_run("PTM: No Data")
    else:
        [completion_date] = x
        doc.tables[1].cell(
            4, 2).paragraphs[-1].add_run(f"PTM: {completion_date.strftime('%m/%d/%Y')}")

    in_ce = False if circuit == "No Data" else circuit_nospace in ce["CE CIRCUIT"].values

    doc.tables[1].cell(4, 2).add_paragraph()
    doc.tables[1].cell(
        4, 2).paragraphs[-1].add_run(f"CE: {'Yes' if in_ce else 'No'}")

    if circuit == "No Data":
        doc.tables[1].cell(4, 2).add_paragraph()
        doc.tables[1].cell(4, 2).paragraphs[-1].add_run(f"2023 Surge: No")
    else:
        x = surge_2023[["WORK_OWNER", "COMPLETION"]][surge_2023["CIRCUIT"] == circuit_nospace].values
        if x.size:
            for record in x:
                [owner, comp_date] = record
                if pd.isna(comp_date):
                    status = "Not Completed"
                elif comp_date > dt.now():
                    status = f"Estimated Completion {comp_date.strftime('%m/%d/%Y')}"
                else:
                    status = f"Completed {comp_date.strftime('%m/%d/%Y')}"
                doc.tables[1].cell(4, 2).add_paragraph()
                doc.tables[1].cell(
                    4, 2).paragraphs[-1].add_run(f"2023 Surge: {status} ({owner})")
        else:
            doc.tables[1].cell(4, 2).add_paragraph()
            doc.tables[1].cell(4, 2).paragraphs[-1].add_run("2023 Surge: No")


    hist_table = doc.add_table(rows=1, cols=7, style="Table Grid")
    colnames = ["Circuit Number", "Outage Type", "Outage Time",
                "Restore Time", "Duration (hh:mm)", "Job ID", "Cause Code"]
    for i, col in enumerate(colnames):
        hist_table.cell(0, i).paragraphs[0].add_run(col).bold = True

    # out_hist_df.drop(columns="latest_circuit", inplace=True)
    for i, (down_date, up_date, duration, circuit, latest_circuit, premise, job_id, cause_code, dur_type) in out_hist_df[out_hist_df["aoe_premise"] == premise].iterrows():
        hist_table.add_row()
        row = hist_table.rows[-1]
        row.cells[0].paragraphs[0].add_run(circuit)
        row.cells[1].paragraphs[0].add_run(dur_type)
        row.cells[2].paragraphs[0].add_run(down_date.strftime("%m/%d/%Y %T"))
        row.cells[3].paragraphs[0].add_run(up_date.strftime("%m/%d/%Y %T"))
        row.cells[4].paragraphs[0].add_run(duration)
        row.cells[5].paragraphs[0].add_run(job_id)
        row.cells[6].paragraphs[0].add_run(cause_code)

    if frequent:
        loc_el = doc.tables[1]._element
        parent_el = loc_el.getparent()
        parent_el.insert(parent_el.index(loc_el)+2, hist_table._element)

    doc.save(f"{cur_dest}{os.sep}Case_Report_{case}_{create_date.replace('/', '')}.docx")


# --- Cleanup

print("Disposing engines...")
adb_engine.dispose()
ora_engine.dispose()
